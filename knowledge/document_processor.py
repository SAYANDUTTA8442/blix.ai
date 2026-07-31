"""
Document Processor — Blix v0.3.2  (Feature 6)

Allows Blix to learn from documents.

Supported formats: PDF, TXT, MD, DOCX, HTML

Pipeline
--------
    Document
       ↓ chunk_document()      — split into overlapping text chunks
       ↓ embed (caller-side)   — EmbeddingStore.encode() per chunk
       ↓ extract_facts()       — LLM/CoT extraction per chunk (or per doc)
       ↓ graph update (caller) — MemoryGraph.upsert_relation() for entities
       ↓ to_memory_entries()   — wrap as MemoryEntry objects for storage

Output
------
``ProcessedDocument`` containing:
    - summary
    - key_findings
    - concepts
    - related_topics
    - chunks (for embedding/indexing)

Python 3.10 compatible. Uses pdfplumber + python-docx (already available);
HTML uses stdlib html.parser to avoid new dependencies.
"""

from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from enum import Enum
from html.parser import HTMLParser
from pathlib import Path
from typing import Optional

from llm.base import LLMProvider
from utils.logger import get_logger

log = get_logger(__name__)


# ---------------------------------------------------------------------------
# Format detection
# ---------------------------------------------------------------------------


class DocumentFormat(str, Enum):
    PDF = "pdf"
    TXT = "txt"
    MD = "md"
    DOCX = "docx"
    HTML = "html"
    UNKNOWN = "unknown"


_EXT_MAP = {
    ".pdf": DocumentFormat.PDF,
    ".txt": DocumentFormat.TXT,
    ".md": DocumentFormat.MD,
    ".markdown": DocumentFormat.MD,
    ".docx": DocumentFormat.DOCX,
    ".html": DocumentFormat.HTML,
    ".htm": DocumentFormat.HTML,
}


def detect_format(path: Path) -> DocumentFormat:
    return _EXT_MAP.get(path.suffix.lower(), DocumentFormat.UNKNOWN)


# ---------------------------------------------------------------------------
# Chunk model
# ---------------------------------------------------------------------------


@dataclass
class DocumentChunk:
    """One chunk of extracted document text, ready for embedding."""

    chunk_id: str
    text: str
    chunk_index: int
    page: Optional[int] = None

    def to_dict(self) -> dict:
        return {
            "chunk_id": self.chunk_id,
            "text": self.text,
            "chunk_index": self.chunk_index,
            "page": self.page,
        }


# ---------------------------------------------------------------------------
# Processed document
# ---------------------------------------------------------------------------


@dataclass
class ProcessedDocument:
    """
    Result of running a document through the processor.

    Fields
    ------
    doc_id:
        Stable id, e.g. derived from filename + timestamp.
    title:
        Document title (filename or extracted heading).
    format:
        ``DocumentFormat``.
    summary:
        High-level natural-language summary.
    key_findings:
        Bullet-style findings (for research papers etc.).
    concepts:
        Notable concepts/terms mentioned.
    related_topics:
        Topic labels for graph/cluster integration.
    entities:
        (label, EntityKind-compatible string) pairs for graph upsert.
    chunks:
        Text chunks for embedding/indexing.
    raw_text_length:
        Total character count of extracted text (for diagnostics).
    """

    doc_id: str
    title: str
    format: DocumentFormat
    summary: str = ""
    key_findings: list[str] = field(default_factory=list)
    concepts: list[str] = field(default_factory=list)
    related_topics: list[str] = field(default_factory=list)
    entities: list[tuple[str, str]] = field(default_factory=list)
    chunks: list[DocumentChunk] = field(default_factory=list)
    raw_text_length: int = 0
    processed_at: str = field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

    def to_dict(self) -> dict:
        return {
            "doc_id": self.doc_id,
            "title": self.title,
            "format": self.format.value,
            "summary": self.summary,
            "key_findings": self.key_findings,
            "concepts": self.concepts,
            "related_topics": self.related_topics,
            "entities": [list(e) for e in self.entities],
            "chunks": [c.to_dict() for c in self.chunks],
            "raw_text_length": self.raw_text_length,
            "processed_at": self.processed_at,
        }


# ---------------------------------------------------------------------------
# Minimal HTML text extractor (stdlib only)
# ---------------------------------------------------------------------------


class _HTMLTextExtractor(HTMLParser):
    _SKIP_TAGS = {"script", "style", "head", "noscript"}

    def __init__(self) -> None:
        super().__init__()
        self._skip_depth = 0
        self._parts: list[str] = []

    def handle_starttag(self, tag: str, attrs) -> None:  # type: ignore[no-untyped-def]
        if tag in self._SKIP_TAGS:
            self._skip_depth += 1
        if tag in ("p", "br", "div", "li", "h1", "h2", "h3", "h4"):
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP_TAGS and self._skip_depth > 0:
            self._skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if self._skip_depth == 0:
            self._parts.append(data)

    def get_text(self) -> str:
        return "".join(self._parts)


# ---------------------------------------------------------------------------
# Extraction prompts
# ---------------------------------------------------------------------------

_DOC_EXTRACT_PROMPT = """\
You are Blix's document analyst. Read the document excerpt below and extract:

1. summary: 2-3 sentence high-level summary
2. key_findings: list of 1-5 concise key findings or claims
3. concepts: list of 3-8 important concepts/terms mentioned
4. related_topics: list of 2-5 short topic labels for categorisation
5. entities: list of [name, type] pairs where type is one of
   person, project, skill, goal, topic, organization

Respond with ONLY a JSON object:
{{
  "summary": "...",
  "key_findings": ["..."],
  "concepts": ["..."],
  "related_topics": ["..."],
  "entities": [["...", "topic"]]
}}

Document excerpt ({title}):
{text}
"""


# ---------------------------------------------------------------------------
# Document Processor
# ---------------------------------------------------------------------------


class DocumentProcessor:
    """
    Extracts text from documents and runs LLM/heuristic analysis.

    Parameters
    ----------
    llm:
        Optional LLM provider for fact/concept extraction. If ``None``,
        heuristic extraction (keyword frequency) is used.
    chunk_size:
        Target characters per chunk.
    chunk_overlap:
        Overlap characters between consecutive chunks.
    max_extract_chars:
        Max characters sent to the LLM for top-level summary/extraction
        (to bound prompt size for large documents).
    """

    def __init__(
        self,
        llm: Optional[LLMProvider] = None,
        chunk_size: int = 1000,
        chunk_overlap: int = 150,
        max_extract_chars: int = 6000,
    ) -> None:
        self._llm = llm
        self._chunk_size = chunk_size
        self._chunk_overlap = chunk_overlap
        self._max_extract = max_extract_chars

    # ------------------------------------------------------------------
    # Entry point
    # ------------------------------------------------------------------

    def process_file(self, path: Path) -> ProcessedDocument:
        """
        Process a document file end-to-end.

        Raises
        ------
        ValueError
            If the format is unsupported or extraction fails.
        """
        fmt = detect_format(path)
        if fmt == DocumentFormat.UNKNOWN:
            raise ValueError(f"Unsupported document format: {path.suffix}")

        text, pages = self._extract_text(path, fmt)
        if not text.strip():
            raise ValueError(f"No text extracted from {path}")

        chunks = self.chunk_document(text, pages)
        doc_id = _slugify(path.stem) + "_" + datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")

        analysis = self._analyse(text, path.stem)

        doc = ProcessedDocument(
            doc_id=doc_id,
            title=path.stem,
            format=fmt,
            summary=analysis.get("summary", ""),
            key_findings=analysis.get("key_findings", []),
            concepts=analysis.get("concepts", []),
            related_topics=analysis.get("related_topics", []),
            entities=[tuple(e) for e in analysis.get("entities", [])],
            chunks=chunks,
            raw_text_length=len(text),
        )
        log.info(
            "DocumentProcessor: processed %s (%s, %d chars, %d chunks)",
            path.name, fmt.value, len(text), len(chunks),
        )
        return doc

    # ------------------------------------------------------------------
    # Text extraction per format
    # ------------------------------------------------------------------

    def _extract_text(
        self, path: Path, fmt: DocumentFormat
    ) -> tuple[str, list[tuple[int, str]]]:
        """
        Returns (full_text, per_page_text) where per_page_text is a list
        of (page_number, text) — empty for formats without page concept.
        """
        if fmt in (DocumentFormat.TXT, DocumentFormat.MD):
            text = path.read_text(encoding="utf-8", errors="replace")
            return text, []

        if fmt == DocumentFormat.HTML:
            raw = path.read_text(encoding="utf-8", errors="replace")
            extractor = _HTMLTextExtractor()
            extractor.feed(raw)
            text = re.sub(r"\n{3,}", "\n\n", extractor.get_text())
            return text.strip(), []

        if fmt == DocumentFormat.DOCX:
            return self._extract_docx(path), []

        if fmt == DocumentFormat.PDF:
            return self._extract_pdf(path)

        raise ValueError(f"No extractor for format {fmt}")

    def _extract_docx(self, path: Path) -> str:
        import docx  # python-docx
        document = docx.Document(str(path))
        parts: list[str] = []
        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    def _extract_pdf(self, path: Path) -> tuple[str, list[tuple[int, str]]]:
        import pdfplumber
        pages: list[tuple[int, str]] = []
        full_parts: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text() or ""
                pages.append((i, page_text))
                full_parts.append(page_text)
        return "\n".join(full_parts), pages

    # ------------------------------------------------------------------
    # Chunking
    # ------------------------------------------------------------------

    def chunk_document(
        self, text: str, pages: Optional[list[tuple[int, str]]] = None
    ) -> list[DocumentChunk]:
        """
        Split ``text`` into overlapping chunks of ~``chunk_size`` characters.

        If ``pages`` is provided (non-empty), chunks are page-aware:
        each page's text is chunked independently and tagged with its
        page number.
        """
        chunks: list[DocumentChunk] = []
        idx = 0

        if pages:
            for page_num, page_text in pages:
                if not page_text.strip():
                    continue
                for piece in self._split_text(page_text):
                    chunks.append(DocumentChunk(
                        chunk_id=f"chunk_{idx}", text=piece, chunk_index=idx, page=page_num,
                    ))
                    idx += 1
        else:
            for piece in self._split_text(text):
                chunks.append(DocumentChunk(chunk_id=f"chunk_{idx}", text=piece, chunk_index=idx))
                idx += 1

        return chunks

    def _split_text(self, text: str) -> list[str]:
        text = text.strip()
        if not text:
            return []
        if len(text) <= self._chunk_size:
            return [text]

        pieces: list[str] = []
        start = 0
        n = len(text)
        step = max(1, self._chunk_size - self._chunk_overlap)
        while start < n:
            end = min(n, start + self._chunk_size)
            # Try to break on whitespace near the boundary
            if end < n:
                boundary = text.rfind(" ", start + int(self._chunk_size * 0.5), end)
                if boundary > start:
                    end = boundary
            pieces.append(text[start:end].strip())
            if end >= n:
                break
            start += step
        return [p for p in pieces if p]

    # ------------------------------------------------------------------
    # Analysis (summary / concepts / entities)
    # ------------------------------------------------------------------

    def _analyse(self, text: str, title: str) -> dict:
        excerpt = text[: self._max_extract]
        if self._llm is not None:
            return self._llm_analyse(excerpt, title)
        return self._heuristic_analyse(excerpt, title)

    def _llm_analyse(self, excerpt: str, title: str) -> dict:
        prompt = _DOC_EXTRACT_PROMPT.format(title=title, text=excerpt)
        try:
            raw = self._llm.generate(prompt).strip()  # type: ignore[union-attr]
            raw = _strip_code_fence(raw)
            data = json.loads(raw)
            return {
                "summary": str(data.get("summary", "")),
                "key_findings": list(data.get("key_findings", [])),
                "concepts": list(data.get("concepts", [])),
                "related_topics": list(data.get("related_topics", [])),
                "entities": [list(e) for e in data.get("entities", []) if len(e) == 2],
            }
        except Exception as exc:
            log.warning("DocumentProcessor: LLM analysis failed (%s); using heuristic.", exc)
            return self._heuristic_analyse(excerpt, title)

    def _heuristic_analyse(self, excerpt: str, title: str) -> dict:
        """Offline fallback: frequency-based concept extraction."""
        words = re.findall(r"[A-Za-z][A-Za-z\-]{3,}", excerpt)
        stop = {
            "this", "that", "with", "from", "have", "been", "were", "they",
            "their", "about", "which", "would", "could", "should", "there",
            "into", "such", "also", "than", "then", "when", "where", "what",
        }
        freq: dict[str, int] = {}
        for w in words:
            lw = w.lower()
            if lw not in stop:
                freq[lw] = freq.get(lw, 0) + 1
        top = sorted(freq.items(), key=lambda kv: -kv[1])[:8]
        concepts = [w for w, _ in top]
        first_sentence = re.split(r"(?<=[.!?])\s", excerpt.strip())[0] if excerpt.strip() else ""
        summary = first_sentence[:300] or f"Document '{title}' processed."
        return {
            "summary": summary,
            "key_findings": [],
            "concepts": concepts,
            "related_topics": concepts[:5],
            "entities": [],
        }


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9_-]", "", text.lower().replace(" ", "_"))


def _strip_code_fence(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        lines = text.splitlines()[1:]
        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines)
    return text.strip()
